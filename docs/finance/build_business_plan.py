"""Build ONE consolidated document: docs/business/WellNest_Business_Plan.docx

Everything in distinct sections: the business plan, the full financial working
(every number derived), the measured control-test validation, the usage/edge-case
model + doctor FAQ, and the v14->v15 change log. All figures mirror
docs/_wellnest_financials.py (Python-verified). Run: python docs/build_business_plan.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x0C, 0x7E, 0xC2); DARK = RGBColor(0x22, 0x2A, 0x35)

# ── Derived, verified figures (from _wellnest_financials.py) ──────────────────
CPN = 0.05
CTS_STD, CTS_PRO = 19.80, 33.00
GM = {"Standard": 79, "Standard + EMR": 86, "Professional": 83, "Professional + EMR": 86}
ARPU, BLEND_CTS, GM_BLEND, GM_STD, BREAKEVEN = 140.10, 24.42, 83, 79, 20
REV_STD = {"Y1": 56_400, "Y2": 225_600, "Y3": 564_000}
REV_BL = {"Y1": 84_060, "Y2": 336_240, "Y3": 840_600}
COGS_STD = {"Y1": 11_880, "Y2": 47_520, "Y3": 118_800}
GP_STD = {"Y1": 44_520, "Y2": 178_080, "Y3": 445_200}
GP_BL = {"Y1": 69_408, "Y2": 277_632, "Y3": 694_080}
OPEX = {"Y1": 55_000, "Y2": 150_000, "Y3": 320_000}
NET_STD = {"Y1": -10_480, "Y2": 28_080, "Y3": 125_200}
NET_BL = {"Y1": 14_408, "Y2": 127_632, "Y3": 374_080}


def M(v): return ("-$" + format(-v, ",.0f")) if v < 0 else ("$" + format(v, ",.0f"))
def h(doc, t, lv=1):
    p = doc.add_heading(t, level=lv)
    for r in p.runs: r.font.color.rgb = BLUE if lv <= 2 else DARK
    return p
def para(doc, t, bold=False, italic=False, size=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p
def bullets(doc, items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        t.rows[0].cells[i].text = ""; t.rows[0].cells[i].paragraphs[0].add_run(x).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    return t
def mono(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(9)
    return p


def build():
    doc = Document(); doc.core_properties.title = "WellNest Business Plan"
    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("WellNest"); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BLUE
    para(doc, "Business Plan & Financials", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "AI-Powered Clinical Documentation, Built by Jamaicans for Caribbean Healthcare", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "CONFIDENTIAL  ·  July 2026  ·  Adrian Tennant & Gary Bryan  ·  (876) 397-6707").alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Single source of truth. Contents: Part A Business Plan · Part B Financial Working "
              "(every number derived) · Part C Measured Cost Validation · Part D Usage Model & Safeguards "
              "· Part E Changes from v14.", italic=True, size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ══ PART A - BUSINESS PLAN ══
    h(doc, "Part A - Business Plan")
    h(doc, "1. Executive Summary", 2)
    para(doc, "Jamaica's healthcare system is at a breaking point: 500+ Cuban clinicians departed, a painful "
              "paper-to-EHR transition adds hours of documentation, physicians spend ~50% of their time on records, "
              "and 70% report burnout. Foreign AI scribes ($900+/mo) lack Patois recognition and offline use.")
    para(doc, "WellNest Scribe turns spoken consultations into structured notes in under 60 seconds, powered by "
              "Cadence - our proprietary pipeline trained on Jamaican Patois. It runs on existing devices, transfers "
              "notes by encrypted QR (no integration), and keeps a clinician in the loop on every note.")
    para(doc, f"Measured unit economics (control test, July 2026): one note costs ~${CPN}; a Standard doctor "
              f"(18 patients/day) costs ~${CTS_STD}/month to serve = ~{GM_STD}% gross margin. Break-even ~{BREAKEVEN} "
              f"clinicians. Seeking US$16,500 non-equity to complete and launch.")
    table(doc, ["Standard-only, conservative", "Year 1", "Year 2", "Year 3"], [
        ["Clinicians", "50", "200", "500"],
        ["Revenue", M(REV_STD["Y1"]), M(REV_STD["Y2"]), M(REV_STD["Y3"])],
        ["Gross margin", f"{GM_STD}%", f"{GM_STD}%", f"{GM_STD}%"],
        ["Net profit (after team)", M(NET_STD["Y1"]), M(NET_STD["Y2"]), M(NET_STD["Y3"])]])

    h(doc, "2. The Problem", 2)
    para(doc, "~0.5 physicians per 1,000 (half the WHO minimum); 500+ Cuban clinicians departed; ~50% of clinician "
              "time on documentation; a paper-based system (handwritten, duplicated, lost); the US$148M HSSP moving "
              "facilities to eCare (13 of 320+ live) creating double-entry burden; a Patois language gap that foreign "
              "speech tools miss ('pressure'=hypertension, 'sugar'=diabetes); and a resulting 70% burnout rate.")
    h(doc, "3. The Solution: WellNest Scribe", 2)
    para(doc, "Under-a-minute structured notes on any browser device; encrypted time-limited QR transfer, no "
              "integration; clinician-reviewed drafts; raw audio not retained; consent captured before recording. "
              "The Cadence engine adds Patois/Creole recognition, clinical structuring, SOAP/narrative/chart "
              "templates, and low-resource-language extensibility. Measured time savings: ~80% on a 5-6 min consult, "
              "~92% on a 20-min consult. Live today: scribe, multilingual recognition, templates, QR sharing, review "
              "UI; in development: lightweight EMR and EMR/EHR integration.")
    h(doc, "4. Market Opportunity", 2)
    para(doc, "Drivers: documentation burden, workforce shortage, the HSSP digital transition, and the Patois gap. "
              "Revenue potential at the US$94 Standard price:")
    table(doc, ["Segment", "Geography", "Physicians", "Annual @ $94"], [
        ["Base market", "EHR-ready (Jamaica)", "500", "$564,000"],
        ["Serviceable", "Jamaica (public+private)", "1,500", "$1,692,000"],
        ["Expansion", "CARICOM (English-speaking)", "11,700", "$13,197,600"]])
    para(doc, "No competitor offers Patois recognition + low-bandwidth resilience + DPA-2020 compliance. Foreign "
              "self-serve tools run ~$30-120/mo; enterprise (Nuance DAX, Abridge, Suki) $300-$1,200. Our moat is "
              "local fit and proprietary Jamaican clinical speech data, not price.")
    h(doc, "5. Business Model & Pricing", 2)
    para(doc, "B2B SaaS, per active clinician per month, billed monthly or annually (annual saves 10%).")
    table(doc, ["Plan", "US$/month", "Gross margin", "Best for"], [
        ["Standard", "94", f"{GM['Standard']}%", "Everyday clinic (~15-20/day)"],
        ["Standard + EMR", "144", f"{GM['Standard + EMR']}%", "Scribe + lightweight EMR"],
        ["Professional", "190", f"{GM['Professional']}%", "High-volume/procedural (~25-40/day)"],
        ["Professional + EMR", "240", f"{GM['Professional + EMR']}%", "Full stack, high volume"],
        ["Institution", "from 76/seat (min 10)", "~75%", "Hospitals / authorities (PO)"]])
    para(doc, "Usage is metered in note-credits: a note up to 20 min = 1 credit; longer recordings cost 1 credit "
              "per 20 min. Standard 500 credits/month (~22 patients/day); Professional 1,100. Doctors are warned "
              "near the limit and never cut off mid-visit (see Part D).")
    h(doc, "6. Go-to-Market", 2)
    para(doc, "Pilot-first: a free pilot at Manchester Health Centre + parallel Ministry engagement, then "
              "public-sector distribution and private direct sales, then CARICOM. Phases: pilot (M1-3), early "
              "adoption (~50 clinicians), island-wide scaling (~200), regional expansion (~500). Channels: direct "
              "sales, referral (1 month free), government partnerships, conferences, resellers, digital marketing.")
    h(doc, "7. Financial Overview", 2)
    para(doc, "All unit economics are measured (control test, July 2026), not estimated. The step-by-step working "
              "is in Part B. Revenue projections:")
    para(doc, "Conservative (every clinician on Standard $94):")
    table(doc, ["", "Year 1", "Year 2", "Year 3"], [
        ["Revenue", M(REV_STD["Y1"]), M(REV_STD["Y2"]), M(REV_STD["Y3"])],
        ["COGS (variable AI)", M(COGS_STD["Y1"]), M(COGS_STD["Y2"]), M(COGS_STD["Y3"])],
        ["Gross profit", M(GP_STD["Y1"]), M(GP_STD["Y2"]), M(GP_STD["Y3"])],
        ["Gross margin", f"{GM_STD}%", f"{GM_STD}%", f"{GM_STD}%"],
        ["Opex (team+platform+marketing)", M(OPEX["Y1"]), M(OPEX["Y2"]), M(OPEX["Y3"])],
        ["Net profit", M(NET_STD["Y1"]), M(NET_STD["Y2"]), M(NET_STD["Y3"])]])
    para(doc, f"Base case (blended tier mix, ARPU ${ARPU}): revenue {M(REV_BL['Y3'])} in Year 3, "
              f"gross profit {M(GP_BL['Y3'])} ({GM_BLEND}%), net profit {M(NET_BL['Y3'])}.")
    para(doc, "Use of the US$16,500 ask: product development & founder stipend (42%), pilot AI compute (11%), "
              "security & compliance (12%), marketing (9%), hosting (7%), onboarding (7%), utilities (5%), "
              "contingency (5%).")
    h(doc, "8. Risk & Mitigation", 2)
    para(doc, "Sales-cycle length (parallel private direct sales); Patois accuracy in noise (clinician review of "
              "every note, weekly tuning, manual fallback); offline reliability (encrypted local caching + "
              "sync-on-reconnect); data privacy (DPA-2020, encryption, audit logging); commoditisation (local-fit "
              "moat); key-person dependence (documented processes, MEC advisory).")
    h(doc, "9. The Team", 2)
    para(doc, "Adrian Tennant & Gary Bryan - Microsoft-certified (AI-102) engineers, final-year Computer Science at "
              "NCU - supported by the Morris Entrepreneurship Centre and clinical advisor Dr. Jerome Smith. Phased "
              "hiring: commission sales (M4), part-time customer success (M10), a Data/ML engineer (Year 2).")
    h(doc, "10. Conclusion", 2)
    para(doc, f"WellNest understands Patois, works offline, and cuts documentation time 80%+. At $94/mo it is "
              f"profitable, with measured gross margins ~{GM_STD}% (Standard) to ~{GM_BLEND}% (blended) and "
              f"break-even at ~{BREAKEVEN} clinicians. Jamaica is our entry point, not our ceiling. Seeking "
              "US$16,500 to complete and launch. WellNest - Saving Time for Those Who Save Lives.")
    doc.add_page_break()

    # ══ PART B - FINANCIAL WORKING ══
    h(doc, "Part B - Financial Working (every number derived)")
    para(doc, "Nothing below is assumed - it is arithmetic from two measured inputs (cost/note and "
              "patients/day). Verified in docs/_wellnest_financials.py.")
    para(doc, "Step 1 - Cost per note (measured control test).", bold=True)
    mono(doc, "5-min : 9,130x$2.50/1M + 618x$15/1M = $0.032 GPT + 30.5s x$0.000164 = $0.005 => $0.037")
    mono(doc, "20-min: 12,667x$2.50/1M + 714x$15/1M = $0.042 GPT + 114s x$0.000164 = $0.019 => $0.061")
    mono(doc, "60-min: 21,882x$2.50/1M + 649x$15/1M = $0.064 GPT + 280s x$0.000164 = $0.046 => $0.110")
    para(doc, f"Typical 12-min note ~$0.045-0.048; we plan conservatively at ${CPN}/note.")
    para(doc, "Step 2 - Cost to serve one clinician / month = notes x cost/note.", bold=True)
    mono(doc, f"Standard    : 18 patients/day x 22 days = 396 notes x ${CPN} = ${CTS_STD}/mo")
    mono(doc, f"Professional: 30 patients/day x 22 days = 660 notes x ${CPN} = ${CTS_PRO}/mo")
    para(doc, "(EMR add-on marginal cost ~$0; the $300/mo platform is a fixed cost carried in Opex.)")
    para(doc, "Step 3 - Gross margin per plan = (price - cost to serve) / price.", bold=True)
    mono(doc, f"Standard ($94): (94 - {CTS_STD})/94 = {GM['Standard']}%   |   Std+EMR ($144): (144 - {CTS_STD})/144 = {GM['Standard + EMR']}%")
    mono(doc, f"Professional ($190): (190 - {CTS_PRO})/190 = {GM['Professional']}%   |   Pro+EMR ($240): (240 - {CTS_PRO})/240 = {GM['Professional + EMR']}%")
    para(doc, "Step 4 - Blended ARPU & cost to serve (mix 50/15/25/10%).", bold=True)
    mono(doc, "ARPU = .50x94 + .15x144 + .25x190 + .10x240 = $140.10")
    mono(doc, f"Blended cost to serve = .65x${CTS_STD} + .35x${CTS_PRO} = ${BLEND_CTS} ; margin {GM_BLEND}%")
    para(doc, "Step 5 - Revenue (verifying $564,000).", bold=True)
    mono(doc, "Conservative Y3 = 500 x $94 x 12 = $564,000   (assumes ALL on the cheapest plan)")
    mono(doc, f"Base case  Y3 = 500 x ${ARPU} x 12 = $840,600   (realistic tier mix)")
    para(doc, "Step 6 - COGS = clinicians x cost to serve x 12.", bold=True)
    mono(doc, f"Std Y3 = 500 x ${CTS_STD} x 12 = $118,800")
    para(doc, "Step 7 - Opex itemised (this is 'employing people').", bold=True)
    mono(doc, "Y3 = founders 2x$40k($80k) + sales team($70k) + customer success($35k) + ML engineer($45k)")
    mono(doc, "     + platform($3.6k) + marketing($50k) + compliance($12k) + misc($24.4k) = $320,000")
    para(doc, "Step 8 - Net profit = Revenue - COGS - Opex.", bold=True)
    mono(doc, f"Std   Y3: $564,000 - $118,800 = $445,200 GP (79%) - $320,000 opex = {M(NET_STD['Y3'])}")
    mono(doc, f"Blend Y3: $840,600 - $146,520 = $694,080 GP (83%) - $320,000 = {M(NET_BL['Y3'])}")
    para(doc, "Step 9 - Break-even & unit economics.", bold=True)
    mono(doc, f"Contribution/Standard clinician = $94 - ${CTS_STD} = $74.20/mo")
    mono(doc, f"Break-even = $1,500/mo fixed / $74.20 = {BREAKEVEN} clinicians")
    mono(doc, "LTV(gross) = $94 x 79% x 30 mo = $2,226 ; CAC $150 ; LTV:CAC 15:1 ; payback 2.0 mo")
    doc.add_page_break()

    # ══ PART C - MEASURED COST VALIDATION ══
    h(doc, "Part C - Measured Cost Validation (control test, 2026-07-11)")
    para(doc, "A live test (real audio, fresh model account) measured the AI cost independently two ways - the "
              "app's own token log and Azure/Modal billing - which agreed to the token.")
    bullets(doc, [
        "Price CONFIRMED at $2.50/$15 per 1M tokens: Azure's gpt-5 daily cost ($0.29) = $2.50/$15 x the day's "
        "93.2K tokens, to the cent. (An assumed $5/$30 rate was ruled out.)",
        "Cost barely scales with length: 5-min $0.037, 20-min $0.061, 60-min $0.110 - 12x the audio, 3x the cost.",
        "Reasoning tokens = 0 on every run; the ~8K-token system prompt is ~90% of GPT cost.",
        "Top optimization lever = prompt caching (cache the static system prompt) - expected to lower GPT cost "
        "further, improving margins.",
        "omniASR/transcription is a rounding error (a few cents per audio-hour, measured).",
    ])
    doc.add_page_break()

    # ══ PART D - USAGE MODEL & SAFEGUARDS ══
    h(doc, "Part D - Usage Model, Safeguards & Doctor FAQ")
    para(doc, "Doctors are capped by note-credits: one patient note up to 20 minutes = 1 credit; a longer "
              "recording costs 1 credit per 20 minutes. This keeps the meter intuitive (a normal note = 1) while "
              "ensuring cost tracks length and cannot be gamed.")
    table(doc, ["Recording length", "Credits"], [
        ["Typed / <=20 min", "1"], ["30 min", "2"], ["1 hour", "3"], ["2 hours", "6"], ["3 hours (auto-stop)", "9"]])
    para(doc, "Safeguards (built):", bold=True)
    bullets(doc, [
        "Hard auto-stop at 3 hours (saves the note, never discards) + on-screen nudge at 90 minutes - a forgotten "
        "recording can't run for hours.",
        "Credit weighting closes the 'many patients in one long recording' loophole (a 4-hour note = 12 credits).",
        "Failed/errored notes are not counted; regenerating a note is free (1 note = 1 credit).",
        "Never cut off mid-visit; overage or a prorated upgrade is available at the cap.",
    ])
    para(doc, "Doctor FAQ:", bold=True)
    bullets(doc, [
        "What counts as a note? One patient encounter (up to 20 min). Longer recordings use a little more.",
        "Is there a time limit? Not a separate clock - time is built into notes (a 1-hour recording = 3 notes).",
        "How many do I get? Standard 500/month, Professional 1,100/month; resets monthly.",
        "Can I record a 2-hour visit? Yes - it's saved and uses 6 notes' worth; auto-stops at 3 hours.",
        "Do edits/regenerations cost extra? No - generating, regenerating, polishing and editing a note = one note.",
    ])
    doc.add_page_break()

    # ══ PART E - CHANGES FROM v14 ══
    h(doc, "Part E - What Changed from the v14 Plan (and why)")
    table(doc, ["Item", "v14", "Now (v15)", "How it's derived"], [
        ["Cost per note", "$0.08-0.15 (est.)", f"${CPN}", "Measured token math (Part B, Step 1)"],
        ["Cost to serve", "$38-42/mo", f"${CTS_STD} Std / ${CTS_PRO} Pro", "notes/mo x $0.05 (Step 2)"],
        ["Gross margin", "55-60%", f"{GM_STD}% Std / {GM_BLEND}% blended", "(price-cost)/price (Step 3)"],
        ["Break-even", "28-30", f"{BREAKEVEN}", "$1,500 fixed / $74.20 (Step 9)"],
        ["Net profit Y3", "$210,000", f"{M(NET_STD['Y3'])} Std / {M(NET_BL['Y3'])} blended", "Rev-COGS-Opex (Step 8)"],
        ["Revenue Y3 ($564k)", "unexplained", "confirmed = 500x$94x12", "= all-Standard case (Step 5)"],
        ["Pricing tiers", "$94/$230/$188/$70.50", "$94/$144/$190/$240/$76-seat", "Current in-app pricing"],
    ])
    para(doc, "The $564,000 was arithmetically correct but assumes every clinician on the cheapest $94 plan; the "
              "realistic blended figure is $840,600.", italic=True)
    doc.add_page_break()

    # ══ PART F - FINANCIAL TERMS EXPLAINED (plain language) ══
    h(doc, "Part F - Financial Terms Explained (plain language)")
    para(doc, "For anyone who isn't a finance person: here is exactly what each term means and how the profit is "
              "built up, step by step. Money comes off in layers - this stack is called a 'P&L' (profit & loss) or "
              "income statement, and it is the standard format every investor expects to see.")
    mono(doc, "  Revenue                 <- all subscription money customers pay")
    mono(doc, "  -  COGS                  <- direct cost to deliver (the AI compute per note)")
    mono(doc, "  =  Gross profit          <- what's left after delivery cost   ('gross margin %')")
    mono(doc, "  -  Operating expenses    <- cost of RUNNING THE COMPANY (salaries, marketing, electricity...)")
    mono(doc, "  =  Net profit            <- the real bottom line you keep      ('net margin %')")
    table(doc, ["Term", "Plain meaning"], [
        ["Revenue", "All the money customers pay you for subscriptions."],
        ["COGS (Cost of Goods Sold)", "The direct cost to deliver the service - here, the AI compute per note (~$19.80/clinician/mo). Nothing else."],
        ["Gross profit / margin", "Revenue minus COGS. 79% means: after paying for the AI we keep 79c of each dollar - BEFORE running the company."],
        ["Operating expenses (Opex)", "The cost of running the company: salaries, sales & marketing, servers, legal, electricity/internet. Not tied to any single note."],
        ["Net profit / margin", "Gross profit minus ALL operating expenses. The real bottom line."],
        ["'After team'", "Shorthand meaning the net-profit row already subtracts salaries (paying the people)."],
    ])
    para(doc, "Why is Year-2 net profit only $28,080 if the margin is 79%?", bold=True)
    para(doc, "Because 79% is the margin on the PRODUCT, not the whole company. Year 2: revenue $225,600 - AI cost "
              "$47,520 = $178,080 gross profit (79%). We then spend $150,000 running the company (mostly hiring "
              "people to grow), leaving $28,080 net. We deliberately reinvest most of the gross profit into the team "
              "early; by Year 3 revenue grows faster than the team, so net profit jumps to $125,200. This is the "
              "normal shape of a growing SaaS business.")
    para(doc, "Analogy: a $10 burger with $2 of ingredients has an 80% gross margin - but after rent, wages and "
              "electricity you might keep $1. The 80% only means the ingredients are cheap; it is not what you pocket.",
         italic=True)
    para(doc, "What 'operating expenses' actually are (itemised):", bold=True)
    table(doc, ["Operating expense", "Year 1", "Year 2", "Year 3"], [
        ["Salaries & wages (founders + hires)", "$38,000", "$103,000", "$230,000"],
        ["Sales & marketing", "$8,000", "$25,000", "$50,000"],
        ["Platform & hosting (servers, DB, GPU)", "$3,600", "$3,600", "$3,600"],
        ["Legal & compliance (DPA, insurance)", "$3,500", "$8,000", "$12,000"],
        ["Office & utilities (electricity, internet, phone)", "$1,900", "$10,400", "$24,400"],
        ["TOTAL operating expenses", "$55,000", "$150,000", "$320,000"]])
    para(doc, "Full income statement (P&L) - the exact numbers, top to bottom:", bold=True)
    nm = {y: round(NET_STD[y] / REV_STD[y] * 100) for y in ("Y1", "Y2", "Y3")}
    table(doc, ["", "Year 1", "Year 2", "Year 3"], [
        ["Revenue", M(REV_STD['Y1']), M(REV_STD['Y2']), M(REV_STD['Y3'])],
        ["-  COGS (AI delivery)", M(COGS_STD['Y1']), M(COGS_STD['Y2']), M(COGS_STD['Y3'])],
        ["=  Gross profit", M(GP_STD['Y1']), M(GP_STD['Y2']), M(GP_STD['Y3'])],
        ["    Gross margin", "79%", "79%", "79%"],
        ["-  Operating expenses", M(OPEX['Y1']), M(OPEX['Y2']), M(OPEX['Y3'])],
        ["=  NET PROFIT", M(NET_STD['Y1']), M(NET_STD['Y2']), M(NET_STD['Y3'])],
        ["    Net margin", f"{nm['Y1']}%", f"{nm['Y2']}%", f"{nm['Y3']}%"]])
    para(doc, "So the two 'margins' to quote an investor: gross margin ~79% (the product is cheap to run) and net "
              "margin rising from negative in Year 1 to ~22% by Year 3 (the company becomes profitable as it scales).",
         italic=True)

    doc.save("docs/business/WellNest_Business_Plan.docx")
    return "docs/business/WellNest_Business_Plan.docx"


if __name__ == "__main__":
    print("Wrote:", build())
